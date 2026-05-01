#!/usr/bin/env python3
"""
Eyeball - 具備原始資料內嵌螢幕截圖的文件分析工具。

將來源文件 (Word、PDF、網頁 URL) 轉換為 PDF，將頁面算繪為圖片，
搜尋引用的文字，醒目提示相符的區域，並組合出一個輸出 
Word 文件，其中分析文字與原始資料螢幕截圖交錯排列。

用法 (由 Copilot CLI 技能呼叫，通常不直接呼叫)：

    python3 eyeball.py build \
        --source <路徑或 URL> \
        --output <輸出 .docx> \
        --sections '[{"heading": "第 1 節", "analysis": "範例分析文字"}]'

    python3 eyeball.py setup-check

    python3 eyeball.py convert --source <檔案 .docx> --output <檔案 .pdf>

    python3 eyeball.py screenshot \
        --source <檔案 .pdf> \
        --anchors '["字詞 1", "字詞 2"]' \
        --page 5 \
        --output screenshot.png
"""

import argparse
import io
import json
import os
import platform
import shutil
import subprocess
import sys
import tempfile

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from PIL import Image, ImageDraw
except ImportError:
    Image = None
    ImageDraw = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None
    Inches = None
    Pt = None
    RGBColor = None


def _resolve_path(path_str):
    """展開使用者提供路徑中的 ~ 和環境變數。"""
    return os.path.expandvars(os.path.expanduser(path_str))


def _check_core_deps():
    """若缺少核心相依性則拋出異常。"""
    missing = []
    if fitz is None:
        missing.append("pymupdf")
    if Image is None:
        missing.append("pillow")
    if Document is None:
        missing.append("python-docx")
    if missing:
        print(f"缺少相依性：{', '.join(missing)}", file=sys.stderr)
        print(f"請執行 setup.sh 或：{sys.executable} -m pip install pymupdf pillow python-docx playwright", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# 文件轉換：來源 -> PDF
# ---------------------------------------------------------------------------

def convert_to_pdf(source_path, output_pdf_path):
    """將文件轉換為 PDF。支援 .docx, .doc, .rtf, .html, .htm。"""
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"找不到來源檔案：{source_path}")

    ext = os.path.splitext(source_path)[1].lower()

    if ext == ".pdf":
        if os.path.abspath(source_path) != os.path.abspath(output_pdf_path):
            shutil.copy2(source_path, output_pdf_path)
        return True

    system = platform.system()

    # 在目前平台上優先嘗試使用 Microsoft Word
    if system == "Darwin" and ext in (".docx", ".doc", ".rtf"):
        if os.path.exists("/Applications/Microsoft Word.app"):
            if _convert_with_word_mac(source_path, output_pdf_path):
                return True

    if system == "Windows" and ext in (".docx", ".doc", ".rtf"):
        if _convert_with_word_windows(source_path, output_pdf_path):
            return True

    # 在任何平台上退而求其次使用 LibreOffice
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice and system == "Windows":
        soffice = _find_libreoffice_windows()
    if soffice and ext in (".docx", ".doc", ".rtf", ".odt", ".html", ".htm"):
        if _convert_with_libreoffice(soffice, source_path, output_pdf_path):
            return True

    raise RuntimeError(
        f"無法將 {ext} 轉換為 PDF。請安裝 Microsoft Word (macOS/Windows) "
        f"或 LibreOffice (任何平台)。"
    )


def _convert_with_word_mac(source_path, output_pdf_path):
    """在 macOS 上透過 AppleScript 使用 Microsoft Word 進行轉換。"""
    source_abs = os.path.abspath(source_path)
    output_abs = os.path.abspath(output_pdf_path)
    # 逸出會破壞 AppleScript 字串插值的字元
    source_safe = source_abs.replace('\\', '\\\\').replace('"', '\\"')
    output_safe = output_abs.replace('\\', '\\\\').replace('"', '\\"')
    script = f'''
    tell application "Microsoft Word"
        open POSIX file "{source_safe}"
        delay 5
        set theDoc to active document
        save as theDoc file name POSIX file "{output_safe}" file format format PDF
        close theDoc saving no
    end tell
    '''
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=120
        )
        return result.returncode == 0 and os.path.exists(output_pdf_path)
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False


def _convert_with_libreoffice(soffice_path, source_path, output_pdf_path):
    """使用 LibreOffice 無頭模式進行轉換。"""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            result = subprocess.run(
                [soffice_path, "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, source_path],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                return False
            basename = os.path.splitext(os.path.basename(source_path))[0] + ".pdf"
            tmp_pdf = os.path.join(tmpdir, basename)
            if os.path.exists(tmp_pdf):
                shutil.move(tmp_pdf, output_pdf_path)
                return True
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
    return False


def _find_libreoffice_windows():
    """在常見的 Windows 安裝位置尋找 LibreOffice。"""
    candidates = []
    for env_var in ("ProgramFiles", "ProgramFiles(x86)"):
        base = os.environ.get(env_var)
        if base:
            candidates.append(os.path.join(base, "LibreOffice", "program", "soffice.exe"))
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _convert_with_word_windows(source_path, output_pdf_path):
    """在 Windows 上透過 win32com 使用 Microsoft Word 進行轉換。"""
    word = None
    doc = None
    try:
        import win32com.client
        source_abs = os.path.abspath(source_path)
        output_abs = os.path.abspath(output_pdf_path)
        os.makedirs(os.path.dirname(output_abs), exist_ok=True)

        # DispatchEx 會建立一個獨立的 Word 程序；如果 DCOM 類別未註冊，則退而求其次使用 Dispatch
        try:
            word = win32com.client.DispatchEx("Word.Application")
        except Exception:
            word = win32com.client.Dispatch("Word.Application")

        word.Visible = False
        word.DisplayAlerts = 0
        try:
            word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
        except Exception:
            pass

        doc = word.Documents.Open(
            FileName=source_abs,
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            NoEncodingDialog=True,
        )
        doc.ExportAsFixedFormat(
            OutputFileName=output_abs,
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
        )
        return os.path.isfile(output_abs)
    except Exception:
        return False
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def render_url_to_pdf(url, output_pdf_path):
    """使用 Playwright 將網頁算繪為 PDF。"""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        raise RuntimeError(
            "需要 Playwright 才能支援網頁 URL。 "
            f"請執行：{sys.executable} -m pip install playwright && "
            f"{sys.executable} -m playwright install chromium"
        )

    with sync_playwright() as p:
        browser = None
        try:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, wait_until="networkidle", timeout=30000)

            # 清除導覽/頁尾元素以獲得更乾淨的輸出
            page.evaluate("""
                document.querySelectorAll(
                    'header, footer, nav, [data-testid="header"], [data-testid="footer"], '
                    + '.site-header, .site-footer, #cookie-banner, .cookie-consent'
                ).forEach(el => el.remove());
            """)

            page.pdf(
                path=output_pdf_path,
                format="Letter",
                print_background=True,
                margin={"top": "0.5in", "bottom": "0.5in",
                        "left": "0.75in", "right": "0.75in"}
            )
        finally:
            if browser is not None:
                browser.close()


# ---------------------------------------------------------------------------
# 螢幕截圖產生
# ---------------------------------------------------------------------------

def screenshot_region(pdf_doc, anchors, target_page=None, target_pages=None,
                      context_padding=40, dpi=200):
    """
    在 PDF 中尋找錨點文字，並將周圍區域擷取為醒目提示的圖片。

    引數：
        pdf_doc: 開啟的 fitz.Document 物件。
        anchors: 搜尋字串列表。裁剪區域會擴展以涵蓋所有字串。
        target_page: 要搜尋的單個頁面 (從 1 開始)。
        target_pages: 要搜尋的多個頁面列表 (結果將垂直拼接)。
        context_padding: 錨點區域上方/下方的額外邊距 (PDF 點數)。
        dpi: 算繪解析度。

    傳回：
        (圖片位元組, 頁面標籤, (寬度, 高度)) 或 (None, None, None)。
    """
    if isinstance(anchors, str):
        anchors = [anchors]

    # 確定要搜尋的頁面
    if target_pages:
        pages = [p - 1 for p in target_pages]
    elif target_page is not None:
        pages = [target_page - 1]
    else:
        pages = list(range(pdf_doc.page_count))

    # 收集跨頁面的符合項
    page_hits = {}
    for pg_idx in pages:
        if pg_idx < 0 or pg_idx >= pdf_doc.page_count:
            continue
        page = pdf_doc[pg_idx]
        hits_on_page = []
        for anchor in anchors:
            found = page.search_for(anchor)
            if found:
                hits_on_page.extend([(anchor, h) for h in found])
        if hits_on_page:
            page_hits[pg_idx] = hits_on_page

    if not page_hits:
        return None, None, None

    zoom = dpi / 72

    # 如果只有單頁，算繪一個區域
    if len(page_hits) == 1:
        pg_idx = list(page_hits.keys())[0]
        img = _render_page_region(pdf_doc, pg_idx, page_hits[pg_idx],
                                   context_padding, zoom)
        img_bytes = _img_to_bytes(img)
        return img_bytes, f"第 {pg_idx + 1} 頁", img.size

    # 多頁：垂直拼接
    images = []
    pages_used = sorted(page_hits.keys())
    for pg_idx in pages_used:
        img = _render_page_region(pdf_doc, pg_idx, page_hits[pg_idx],
                                   context_padding, zoom)
        images.append(img)

    stitched = _stitch_vertical(images)
    img_bytes = _img_to_bytes(stitched)

    if len(pages_used) > 1:
        page_nums = ", ".join(str(p + 1) for p in pages_used)
        page_label = f"第 {page_nums} 頁"
    else:
        page_label = f"第 {pages_used[0]+1} 頁"

    return img_bytes, page_label, stitched.size


def _render_page_region(pdf_doc, pg_idx, hits_with_anchors, context_padding, zoom):
    """算繪 PDF 頁面的裁剪區域，並醒目提示錨點文字。"""
    page = pdf_doc[pg_idx]
    page_rect = page.rect

    all_rects = [h for _, h in hits_with_anchors]
    min_y = min(r.y0 for r in all_rects)
    max_y = max(r.y1 for r in all_rects)

    crop_rect = fitz.Rect(
        page_rect.x0 + 20,
        max(page_rect.y0, min_y - context_padding),
        page_rect.x1 - 20,
        min(page_rect.y1, max_y + context_padding)
    )

    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, clip=crop_rect)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # 醒目提示每個錨點符合項
    draw = ImageDraw.Draw(img, "RGBA")
    pad = max(2, round(2 * zoom))
    for anchor, rect in hits_with_anchors:
        if rect.y0 >= crop_rect.y0 - 5 and rect.y1 <= crop_rect.y1 + 5:
            x0 = (rect.x0 - crop_rect.x0) * zoom
            y0 = (rect.y0 - crop_rect.y0) * zoom
            x1 = (rect.x1 - crop_rect.x0) * zoom
            y1 = (rect.y1 - crop_rect.y0) * zoom
            draw.rectangle([x0-pad, y0-pad, x1+pad, y1+pad], fill=(255, 255, 0, 100))

    # 邊框
    ImageDraw.Draw(img).rectangle(
        [0, 0, img.width - 1, img.height - 1],
        outline=(160, 160, 160), width=2
    )

    return img


def _stitch_vertical(images, gap=4):
    """垂直拼接多張圖片，並在它們之間留出小間隙。"""
    total_height = sum(img.height for img in images) + gap * (len(images) - 1)
    max_width = max(img.width for img in images)
    stitched = Image.new("RGB", (max_width, total_height), (255, 255, 255))
    y = 0
    for img in images:
        stitched.paste(img, (0, y))
        y += img.height + gap
    ImageDraw.Draw(stitched).rectangle(
        [0, 0, stitched.width - 1, stitched.height - 1],
        outline=(160, 160, 160), width=2
    )
    return stitched


def _img_to_bytes(img):
    """將 PIL Image 轉換為 PNG BytesIO 緩衝區 (類檔案物件)。"""
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 輸出文件組合
# ---------------------------------------------------------------------------

def build_analysis_doc(pdf_doc, sections, output_path, title=None, subtitle=None,
                       source_label=None, dpi=200):
    """
    建置包含分析章節和原始資料內嵌螢幕截圖的 Word 文件。

    引數：
        pdf_doc: 開啟的 fitz.Document (來源，已轉換為 PDF)。
        sections: 字典列表，每個字典包含：
            - heading (str): 章節標題
            - analysis (str): 分析文字
            - anchors (list[str]): 要搜尋並醒目提示的原始資料逐字片語
            - target_page (int, 選填): 要搜尋的單一頁面 (從 1 開始)
            - target_pages (list[int], 選填): 要搜尋的多個頁面
            - context_padding (int, 選填): 額外的 PDF 點數邊距 (預設 40)
        output_path: 輸出 .docx 檔案的儲存路徑。
        title: 文件標題。
        subtitle: 文件副標題。
        source_label: 來源標籤 (例如檔名或 URL)。
        dpi: 螢幕截圖解析度。
    """
    doc = Document()

    # 樣式
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 標題
    if title:
        doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph("")

    # 章節
    for i, section in enumerate(sections):
        heading = section.get("heading", f"第 {i+1} 節")
        analysis = section.get("analysis", "")
        anchors = section.get("anchors", [])
        target_page = section.get("target_page")
        target_pages = section.get("target_pages")
        padding = section.get("context_padding", 40)

        doc.add_heading(heading, level=2)
        doc.add_paragraph(analysis)

        if anchors:
            img_bytes, page_label, size = screenshot_region(
                pdf_doc, anchors,
                target_page=target_page,
                target_pages=target_pages,
                context_padding=padding,
                dpi=dpi
            )

            if img_bytes:
                # 來源標籤
                p = doc.add_paragraph()
                anchor_text = ", ".join(f'"{a}"' for a in anchors[:3])
                if len(anchors) > 3:
                    anchor_text += f" (以及另外 {len(anchors)-3} 個)"
                label = f"[來源：{source_label or '文件'}，{page_label}"
                label += f" -- 醒目提示：{anchor_text}]"
                run = p.add_run(label)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(120, 120, 120)
                run.font.italic = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)

                # 螢幕截圖
                doc.add_picture(img_bytes, width=Inches(5.8))
                doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            else:
                # 找不到錨點
                p = doc.add_paragraph()
                run = p.add_run(
                    f"[無法提供螢幕截圖：在來源文件中找不到 "
                    f"{', '.join(repr(a) for a in anchors)}]"
                )
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(180, 50, 50)

    # 頁尾備註
    doc.add_paragraph("")
    note = doc.add_paragraph()
    run = note.add_run(
        "由 Eyeball 產生。每張螢幕截圖均擷取自來源文件，引用的文字以黃色反白。 "
        "螢幕截圖會動態調整大小，以涵蓋分析中參考的完整文字範圍。 "
        "請檢閱醒目提示的原始資料以驗證每項斷言。"
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI 指令
# ---------------------------------------------------------------------------

def cmd_setup_check():
    """檢查所有相依性是否可用。"""
    checks = {
        "PyMuPDF": False,
        "Pillow": False,
        "python-docx": False,
        "Playwright": False,
        "Chromium 瀏覽器": False,
        "Word (macOS)": False,
        "Word (Windows)": False,
        "LibreOffice": False,
    }

    try:
        import fitz
        checks["PyMuPDF"] = True
    except ImportError:
        pass

    try:
        from PIL import Image
        checks["Pillow"] = True
    except ImportError:
        pass

    try:
        from docx import Document
        checks["python-docx"] = True
    except ImportError:
        pass

    try:
        from playwright.sync_api import sync_playwright
        checks["Playwright"] = True
    except ImportError:
        pass

    # 在所有平台上檢查 Chromium
    pw_cache_candidates = []
    system = platform.system()
    if system == "Darwin":
        pw_cache_candidates.append(os.path.expanduser("~/Library/Caches/ms-playwright"))
    if system == "Windows":
        local_app_data = os.environ.get("LOCALAPPDATA", "")
        if local_app_data:
            pw_cache_candidates.append(os.path.join(local_app_data, "ms-playwright"))
    pw_cache_candidates.append(os.path.expanduser("~/.cache/ms-playwright"))
    # 尊重 PLAYWRIGHT_BROWSERS_PATH
    custom_pw = os.environ.get("PLAYWRIGHT_BROWSERS_PATH")
    if custom_pw and custom_pw != "0":
        pw_cache_candidates.insert(0, custom_pw)
    for pw_cache in pw_cache_candidates:
        if os.path.isdir(pw_cache) and any(
            d.startswith("chromium") for d in os.listdir(pw_cache)
        ):
            checks["Chromium 瀏覽器"] = True
            break

    # 檢查轉換器 -- 僅限登錄表/檔案系統，絕不啟動 Word
    if system == "Darwin" and os.path.exists("/Applications/Microsoft Word.app"):
        checks["Word (macOS)"] = True

    if system == "Windows":
        try:
            import winreg
            word_reg_paths = [
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE"),
                (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE"),
                (winreg.HKEY_CLASSES_ROOT, r"Word.Application"),
            ]
            for hive, subkey in word_reg_paths:
                try:
                    winreg.OpenKey(hive, subkey)
                    checks["Word (Windows)"] = True
                    break
                except OSError:
                    pass
        except ImportError:
            pass
        # 檢查是否可用 pywin32 進行 Word 自動化
        if checks["Word (Windows)"]:
            try:
                import win32com.client  # noqa: F401
            except ImportError:
                checks["Word (Windows)"] = False
                print("  備註：已發現 Microsoft Word，但未安裝 pywin32。")
                print(f"  請執行：{sys.executable} -m pip install pywin32")

    if shutil.which("libreoffice") or shutil.which("soffice"):
        checks["LibreOffice"] = True
    elif system == "Windows":
        if _find_libreoffice_windows():
            checks["LibreOffice"] = True

    print("Eyeball 相依性檢查：")
    all_core = True
    for name, ok in checks.items():
        status = "正常 (OK)" if ok else "缺少 (MISSING)"
        marker = "+" if ok else "-"
        print(f"  [{marker}] {name}：{status}")
        if name in ("PyMuPDF", "Pillow", "python-docx") and not ok:
            all_core = False

    has_converter = checks["Word (macOS)"] or checks["Word (Windows)"] or checks["LibreOffice"]
    has_web = checks["Playwright"] and checks["Chromium 瀏覽器"]

    print("")
    print("來源支援：")
    print(f"  PDF 檔案：   {'就緒 (Ready)' if all_core else '需要執行：pip3 install pymupdf pillow python-docx'}")
    print(f"  Word 文件：   {'就緒 (Ready)' if has_converter else '需要：Microsoft Word 或 LibreOffice'}")
    print(f"  網頁 URL：    {'就緒 (Ready)' if has_web else '需要執行：pip3 install playwright && python3 -m playwright install chromium'}")

    return 0 if all_core else 1


def cmd_convert(args):
    """將文件轉換為 PDF。"""
    source = _resolve_path(args.source)
    output = _resolve_path(args.output)

    if source.startswith(("http://", "https://")):
        render_url_to_pdf(source, output)
    else:
        convert_to_pdf(source, output)

    print(f"已轉換：{output} ({os.path.getsize(output)} 位元組)")


def cmd_screenshot(args):
    """從 PDF 產生單張螢幕截圖。"""
    _check_core_deps()
    source = _resolve_path(args.source)

    if not os.path.isfile(source):
        print(f"找不到來源檔案：{source}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(source)[1].lower()
    if ext != ".pdf":
        print(f"來源必須是 PDF 檔案 (目前為 {ext})。 "
              f"請先使用 'convert' 轉換其他格式。", file=sys.stderr)
        sys.exit(1)

    anchors = json.loads(args.anchors)
    target_page = args.page
    padding = args.padding
    dpi = args.dpi

    pdf_doc = fitz.open(source)
    try:
        img_bytes, page_label, size = screenshot_region(
            pdf_doc, anchors,
            target_page=target_page,
            context_padding=padding,
            dpi=dpi
        )
    finally:
        pdf_doc.close()

    if img_bytes:
        output = _resolve_path(args.output)
        with open(output, "wb") as f:
            f.write(img_bytes.getvalue())
        print(f"螢幕截圖已儲存：{output} ({size[0]}x{size[1]} 像素，{page_label})")
    else:
        print(f"找不到相符項：{anchors}", file=sys.stderr)
        sys.exit(1)


def cmd_build(args):
    """建置完整分析文件。"""
    _check_core_deps()
    source = _resolve_path(args.source)
    output = _resolve_path(args.output)
    sections = json.loads(args.sections)
    title = args.title
    subtitle = args.subtitle
    dpi = args.dpi

    if not source.startswith(("http://", "https://")) and not os.path.isfile(source):
        print(f"找不到來源檔案：{source}", file=sys.stderr)
        sys.exit(1)

    # 確定來源類型並轉換為 PDF
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_pdf = tmp.name

    pdf_doc = None
    try:
        if source.startswith(("http://", "https://")):
            render_url_to_pdf(source, tmp_pdf)
            source_label = source
        elif source.lower().endswith(".pdf"):
            shutil.copy2(source, tmp_pdf)
            source_label = os.path.basename(source)
        else:
            convert_to_pdf(source, tmp_pdf)
            source_label = os.path.basename(source)

        pdf_doc = fitz.open(tmp_pdf)
        build_analysis_doc(
            pdf_doc, sections, output,
            title=title, subtitle=subtitle,
            source_label=source_label,
            dpi=dpi
        )

        size_kb = os.path.getsize(output) / 1024
        print(f"分析已儲存：{output} ({size_kb:.0f} KB)")

    finally:
        if pdf_doc is not None:
            pdf_doc.close()
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)


def cmd_extract_text(args):
    """從來源文件提取文字 (供 AI 在撰寫分析前閱讀)。"""
    _check_core_deps()
    source = _resolve_path(args.source)

    if not source.startswith(("http://", "https://")) and not os.path.isfile(source):
        print(f"找不到來源檔案：{source}", file=sys.stderr)
        sys.exit(1)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_pdf = tmp.name

    pdf_doc = None
    try:
        if source.startswith(("http://", "https://")):
            render_url_to_pdf(source, tmp_pdf)
        elif source.lower().endswith(".pdf"):
            shutil.copy2(source, tmp_pdf)
        else:
            convert_to_pdf(source, tmp_pdf)

        pdf_doc = fitz.open(tmp_pdf)
        for i in range(pdf_doc.page_count):
            text = pdf_doc[i].get_text()
            print(f"\n[第 {i+1} 頁]")
            print(text)

    finally:
        if pdf_doc is not None:
            pdf_doc.close()
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)


def main():
    parser = argparse.ArgumentParser(
        description="Eyeball：具備原始資料內嵌螢幕截圖的文件分析工具"
    )
    sub = parser.add_subparsers(dest="command")

    # setup-check
    sub.add_parser("setup-check", help="檢查相依性")

    # convert
    p_conv = sub.add_parser("convert", help="將文件轉換為 PDF")
    p_conv.add_argument("--source", required=True)
    p_conv.add_argument("--output", required=True)

    # screenshot
    p_ss = sub.add_parser("screenshot", help="從 PDF 產生一張螢幕截圖")
    p_ss.add_argument("--source", required=True, help="PDF 檔案路徑")
    p_ss.add_argument("--anchors", required=True, help="搜尋詞的 JSON 陣列")
    p_ss.add_argument("--page", type=int, help="目標頁面 (從 1 開始)")
    p_ss.add_argument("--padding", type=int, default=40)
    p_ss.add_argument("--dpi", type=int, default=200)
    p_ss.add_argument("--output", required=True, help="輸出 PNG 路徑")

    # build
    p_build = sub.add_parser("build", help="建置分析文件")
    p_build.add_argument("--source", required=True,
                         help="來源文件路徑或 URL")
    p_build.add_argument("--output", required=True,
                         help="輸出 .docx 路徑")
    p_build.add_argument("--sections", required=True,
                         help="章節物件的 JSON 陣列")
    p_build.add_argument("--title", help="文件標題")
    p_build.add_argument("--subtitle", help="文件副標題")
    p_build.add_argument("--dpi", type=int, default=200)

    # extract-text
    p_text = sub.add_parser("extract-text",
                            help="從文件提取文字 (供 AI 分析使用)")
    p_text.add_argument("--source", required=True)

    args = parser.parse_args()

    if args.command == "setup-check":
        sys.exit(cmd_setup_check())
    elif args.command == "convert":
        cmd_convert(args)
    elif args.command == "screenshot":
        cmd_screenshot(args)
    elif args.command == "build":
        cmd_build(args)
    elif args.command == "extract-text":
        cmd_extract_text(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
